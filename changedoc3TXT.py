# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
import os
import gender_detector
import docxToPdf
import email_sender
import clean_directory
import space_remover
from datetime import date

# print("Initializing donor document...")
doc = Document('donor3.docx')  # read default document
listOfLines = list()
with open('Untitled.txt', 'r', encoding='utf8') as txt:
    for line in txt:
        listOfLines.append(line.strip())

#lines = txt.readlines()
# for a in lines: u(a)

print(listOfLines)
print('Initializing donor document "donor3.docx"... OK')
print('Initializing text "Untitled.txt"... OK')
print("""
              _,     _   _    ,_
           o888P     Y8o8Y     Y888o.
         d88888      88888      88888b
       ,8888888b_  _d88888b_  _d8888888,
       888888888888888888888888888888888
       888888888888888888888888888888888
        Y8888P"Y888P"Y888P-Y888P"Y88888'
         Y888   '8'   Y8P   '8'   888Y
          '8o          V          o8'
             `                   `
""")

d = "d"  # d is always d, right?
vienna = {
    1010: "Innere Stadt",
    1020: 'Leopoldstadt',
    1030: 'Landstraße',
    1040: 'Wieden',
    1050: 'Margareten',
    1060: 'Mariahilf',
    1070: 'Neubau',
    1080: 'Josefstadt',
    1090: 'Alsergrund',
    1100: 'Favoriten',
    1110: 'Simmering',
    1120: 'Meidling',
    1130: 'Hietzing',
    1140: 'Penzing',
    1150: 'Rudolfsheim-Fünfhaus',
    1160: 'Ottakring',
    1170: 'Hernals',
    1180: 'Währing',
    1190: 'Döbling',
    1200: 'Brigittenau',
    1210: 'Floridsdorf',
    1220: 'Donaustadt',
    1230: 'Liesing',
    1300: 'Schwechat',
    1400: 'Vienna International Centre'
}

gruss_receiver = listOfLines[0]
Obj = space_remover.SpaceRemover(gruss_receiver)
gruss_receiver = Obj.remove()
mail_receiver = gruss_receiver
if gruss_receiver == d:
    gruss_receiver = "Damen und Herren"
    mail_receiver = "Personalabteilung"
print("Okay, sending to ", gruss_receiver)

job = listOfLines[1]
if job == d:
    job = "Mitarbeiter"
Obj = space_remover.SpaceRemover(job)
job = Obj.remove()
print("Okay, your possible job is ", job)

firma_receiver = listOfLines[2]
if firma_receiver == d:
    firma_receiver = "Ihrer Firma"
Obj = space_remover.SpaceRemover(firma_receiver)
firma_receiver = Obj.remove()
print("Okay, let's write ", firma_receiver)

address = listOfLines[3]
if address == d:
    address = " "
else:
    Obj = space_remover.SpaceRemover(address)
    address = Obj.remove()
print("Okay, address is ", address)

zipcode = listOfLines[4]
Obj = space_remover.SpaceRemover(zipcode)
zipcode = Obj.remove()
if job == d:
    zipcode = 1010

try:
    k = int(zipcode)
    if k in vienna:
        zipcode = str(zipcode) + " " + str(vienna[k])
except:
    print("Can't resolve ", zipcode, "as key for dict")

print("Okay, ZIP is ", zipcode)

paragraph1 = doc.paragraphs[1].text
paragraph2 = doc.paragraphs[2].text
paragraph3 = doc.paragraphs[3].text
paragraph4 = doc.paragraphs[4].text

print("Parsing paragraphs... OK")

print(gruss_receiver[0])
gruss_receiver_gender = d

# Detecting a gender of a receiver for applying into Word document
Receiver = gender_detector.GenderDetector(gruss_receiver)
temp = Receiver.detect_gender()
print(temp)

paragraph1 = paragraph1.replace("%%firma%%", firma_receiver)
print("Putting firma name '" + firma_receiver + "', OK")

paragraph1 = paragraph1.replace("%%receiver1%%", mail_receiver)
print("Putting name '" + mail_receiver + "', OK")

paragraph1 = paragraph1.replace("%%address%%", address)
print("Putting address '" + address + "', OK")

paragraph1 = paragraph1.replace("%%zip%%", zipcode)
print("Putting zip '" + zipcode + "', OK")

today = date.today().strftime("%d.%m.%Y")
paragraph2 = paragraph2.replace("%%date%%", today)
print("Putting today's date, OK")

paragraph3 = paragraph3.replace("%%job%%", job)
print("Putting job name '" + job + "', OK")

if temp == "Male":
    paragraph4 = paragraph4.replace("%%receiver2%%", "r " + gruss_receiver)
    print("Putting male name '" + gruss_receiver + "', OK")

elif temp == "Female":
    paragraph4 = paragraph4.replace("%%receiver2%%", " " + gruss_receiver)
    print("Putting female '" + gruss_receiver + "', OK")
else:
    paragraph4 = paragraph4.replace("%%receiver2%%", " " + gruss_receiver)
    print("Putting default name '" + gruss_receiver + "', OK")

doc.paragraphs[1].text = paragraph1
doc.paragraphs[2].text = paragraph2
doc.paragraphs[3].text = paragraph3
doc.paragraphs[4].text = paragraph4
print("Pasting text... OK")

# doc.paragraphs[0].style = "bewerbung_default_paragraph"
doc.paragraphs[1].style = "bewerbung_receiver"
doc.paragraphs[2].style = "bewerbung_date_right"
doc.paragraphs[3].style = "bewerbung_header"
doc.paragraphs[4].style = "bewerbung_default_paragraph"
print("Applying styles... OK")

# saving_name = job + ".docx"
saving_name = "Anschreiben als " + job + " in " + firma_receiver
if gruss_receiver_gender != d:
    saving_name = saving_name + " - " + gruss_receiver
saving_name_docx = saving_name + ".docx"
saving_name_pdf = saving_name + ".pdf"
doc.save(saving_name_docx)

#####CONVERTING DOCX INTO PDF########

pdf = docxToPdf.Converter(saving_name_pdf, saving_name_docx)
pdf.convert()
# saving_name_pdf = saving_name_pdf.replace(" fuer ", " für ")

print("""
 ..............
 .  *      *  .
[.      '     .]
 .            .
 .     ~~     .
 ..............
""")
####SEND BY EMAIL#######

receiver_email = ""

temp01_email = listOfLines[5]
if temp01_email == "Y" or temp01_email == "y":
    receiver_email = input(">> EMAIL of the receiver: ")
    email = email_sender.EmailSender(gruss_receiver, receiver_email, job, gruss_receiver_gender, saving_name_pdf)
    email.send_email()
elif temp01_email == "N" or temp01_email == "n":
    print("Not sending an e-mail. Goodbye!")
else:
    print("Something went wrong")




os.remove(saving_name_docx)
cleaner = clean_directory.Mover(saving_name_docx, saving_name_pdf, firma_receiver, receiver_email)
cleaner.clean()
