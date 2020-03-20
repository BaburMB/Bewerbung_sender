from docx import Document
from docx.shared import Pt
from docx.shared import Inches
import os
import gender_detector
import docxToPdf
import email_sender
import clean_directory
import space_remover


print("Initializing donor document...")
doc = Document('donor.docx')                    #read default document
print("Initializing donor document... OK")


d = "d"                                         #d is always d, right?


gruss_receiver = input(">> WHOM are you sending? ")
if gruss_receiver == d:
    gruss_receiver = "Damen und Herren"

Obj = space_remover.SpaceRemover(gruss_receiver)
gruss_receiver = Obj.remove()

print("Okay, sending to ", gruss_receiver)

job = input(">> What kind of JOB is that? ")
if job == d:
    job = "Mitarbeiter"

Obj = space_remover.SpaceRemover(job)
job = Obj.remove()

print("Okay, your possible job is ", job)

firma_receiver = input(">> What FIRMA is that? ")
if firma_receiver == d:
    firma_receiver = "Ihrer Firma"

Obj = space_remover.SpaceRemover(firma_receiver)
firma_receiver = Obj.remove()

print("Okay, let's write ", firma_receiver)

print("Parsing paragraph...")
paragraph = doc.paragraphs[2].text
print("Parsing paragraph... OK")

print(gruss_receiver[0])
gruss_receiver_gender = d

# Detecting a gender of a receiver for applying into Word document
Receiver = gender_detector.GenderDetector(gruss_receiver)
temp = Receiver.detect_gender()
print(temp)


if temp == "Male":
    paragraph = paragraph.replace("%%receiver%%", "r " + gruss_receiver)
    print("Putting male name " + gruss_receiver + ", OK")

elif temp == "Female":
    paragraph = paragraph.replace("%%receiver%%", " " + gruss_receiver)
    print("Putting female " + gruss_receiver + ", OK")
else:
    paragraph = paragraph.replace("%%receiver%%", " " + gruss_receiver)
    print("Putting default name " + gruss_receiver + ", OK")

paragraph = paragraph.replace("%%firma%%", firma_receiver)
print("Putting firma name " + firma_receiver + ", OK")
paragraph = paragraph.replace("%%job%%", job)
print("Putting job name " + job + ", OK")
doc.paragraphs[2].text = paragraph
print("Pasting text... OK")
doc.paragraphs[2].style = "bewerbung_default_paragraph"
print("Applying style... OK")

# saving_name = job + ".docx"
saving_name = "Anschreiben als " + job + " in " + firma_receiver
if gruss_receiver_gender != d:
    saving_name = saving_name + " - " + gruss_receiver
saving_name_docx = saving_name + ".docx"
saving_name_pdf = saving_name + ".pdf"
doc.save(saving_name_docx)

#####CONVERTING DOCX INTO PDF########

pdf = docxToPdf.Converter(saving_name_pdf, saving_name_docx)
pdf.convert()
# saving_name_pdf = saving_name_pdf.replace(" fuer ", " für ")


####SEND BY EMAIL#######

receiver_email = ""
while True:
    temp01_email = input("Send E-Mail? Y/N ")
    if temp01_email == "Y" or temp01_email == "y":
        receiver_email = input(">> EMAIL of the receiver: ")
        email = email_sender.EmailSender(gruss_receiver, receiver_email, job, gruss_receiver_gender, saving_name_pdf)
        email.send_email()
        break
    elif temp01_email == "N" or temp01_email == "n":
        print("Not sending an e-mail. Goodbye!")
        break
    else:
        continue

os.remove(saving_name_docx)
cleaner = clean_directory.Mover(saving_name_docx, saving_name_pdf, firma_receiver, receiver_email)
cleaner.clean()

