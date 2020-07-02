# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
import os
import gender_detector
import docxToPdf
import email_sender
import clean_directory
import space_remover as sr
from datetime import date


class DocAbstract:

    def __init__(self):
        pass

    def run(self):

        doc_english = 'donor3ENG.docx'
        doc_german = 'donor3.docx'
        # Initializing input lines
        # 1. Language:
        #       "e" - for English
        #       "g" - for German
        # 2. Receiver:
        #       "d" - for default
        # 3. Job position:
        #       "d" - for Mitarbeiter
        # 4. Company:
        #       "d" - for default
        # 5. Address
        # 6. ZIP code
        # 7. E-mail sending. Under construction, leave as n
        #
        #
        vienna = {
            1010: "Innere Stadt",
            1020: 'Leopoldstadt',
            1030: 'Landstraße',
            1040: 'Wieden',
            1050: 'Margareten',
            1060: 'Mariahilf',
            1070: 'Neubau',
            1080: 'Josefstadt',
            1090: 'Alsergrund',
            1100: 'Favoriten',
            1110: 'Simmering',
            1120: 'Meidling',
            1130: 'Hietzing',
            1140: 'Penzing',
            1150: 'Rudolfsheim-Fünfhaus',
            1160: 'Ottakring',
            1170: 'Hernals',
            1180: 'Währing',
            1190: 'Döbling',
            1200: 'Brigittenau',
            1210: 'Floridsdorf',
            1220: 'Donaustadt',
            1230: 'Liesing',
            1300: 'Schwechat',
            1400: 'Vienna International Centre'
        }

        listOfLines = list()
        with open('Untitled.txt', 'r', encoding='utf8') as txt:
            for line in txt:
                listOfLines.append(line.strip())

        lanGerman = False if listOfLines[6] == "e" else True
        doc = Document(doc_german) if lanGerman else Document(doc_english)

        # initializing every object
        gruss_receiver = sr.SpaceRemover(listOfLines[0]).remove()
        mail_receiver = gruss_receiver
        if gruss_receiver == d:
            if lanGerman:
                gruss_receiver = "Damen und Herren"
                mail_receiver = "Personalabteilung"
            else:
                gruss_receiver = "Sir or Madam"
                mail_receiver = "Human Resources"

        # JOB

        firma_receiver = sr.SpaceRemover(listOfLines[2]).remove()
        if firma_receiver == d:
            firma_receiver = "Ihrer Firma" if lanGerman else "Your company"

        address = sr.SpaceRemover(listOfLines[3]).remove()
        if firma_receiver == d:
            firma_receiver = " "

        zipcode = sr.SpaceRemover(listOfLines[4]).remove()
        if job == d:
            zipcode = 1010
        try:
            k = int(zipcode)
            if k in vienna:
                zipcode = str(zipcode) + " " + str(vienna[k])
        except: print("Can't resolve ", zipcode, "as key for dict")

        paragraph1 = doc.paragraphs[1].text
        paragraph2 = doc.paragraphs[2].text
        paragraph3 = doc.paragraphs[3].text
        paragraph4 = doc.paragraphs[4].text
        paragraph13 = doc.paragraphs[13].text

        #logic

        doc.paragraphs[1].text = paragraph1
        doc.paragraphs[2].text = paragraph2
        doc.paragraphs[3].text = paragraph3
        doc.paragraphs[4].text = paragraph4
        doc.paragraphs[13].text = paragraph13
        print("Pasting text... OK")

        doc.paragraphs[1].style = "bewerbung_receiver"
        doc.paragraphs[2].style = "bewerbung_date_right"
        doc.paragraphs[3].style = "bewerbung_header"
        doc.paragraphs[4].style = "bewerbung_default_paragraph"
        doc.paragraphs[13].style = "bewerbung_default_paragraph"
        print("Applying styles... OK")

