from docx import Document
from docx.shared import Inches

document = Document()

default_gruss = "Damen und Herren"
gruss_receiver_gender = "default"
d = "d"


gruss_receiver = input("WHOM are you sending? ")
if gruss_receiver == d: gruss_receiver = "Damen und Herren"
print("Okay, sending to ", gruss_receiver)
job = input("What kind of JOB is that? ")
if job == d: job = "Mitarbeiter"
print("Okay, your possible job is ", job)
firma_receiver = input("What FIRMA is that? ")
if firma_receiver == d: firma_receiver = "Ihrer Firma"
print("Okay, let's write ", firma_receiver)


# if gruss_receiver[0] == " ":
#     print("You have a space!")
# else:
#     print("Checking gender...")

print(gruss_receiver[0])

if gruss_receiver[0] == "F" or gruss_receiver[0] == "f":
    gruss_receiver_gender = "Female"
elif gruss_receiver[0] == "H" or gruss_receiver[0] == "h":
    gruss_receiver_gender = "Male"
else:
    print("Couldn't detect gender with ", gruss_receiver[0], ", will use default gender for Grüß")
print(gruss_receiver_gender)

# document.add_heading('Document Title', 0)
# Sehr geerte...
sehr = 'Sehr geehrte'


p = document.add_paragraph(sehr)

if gruss_receiver_gender == "Male":
    p.add_run('r ')
    p.add_run(gruss_receiver)

elif gruss_receiver_gender == "Female":
    p.add_run(' ')
    p.add_run(gruss_receiver)
else:
    p.add_run(' ')
    p.add_run(default_gruss)

p.add_run(',')


# p.add_run('italic.').italic = True
#Adding last parargraph
print("Adding last paragraph")
document.add_paragraph('''
Gerne bringe ich meine Berufserfahrung und umfassende Kompetenz in Ihr Unternehmen ein. Nicht nur in der von Ihnen ausgeschriebene Stelle, sondern im gesamten Unternehmen und Marktfeld sehe ich enorme Chancen, meine Qualifikationen für Sie gewinnbringend einzusetzen und beispielsweise Kosten zu sparen und Produkte zu optimieren. Mit mir erhalten Sie einen Mitarbeiter, der hoch motiviert, selbstständig und kreativ nach neuen Lösungen sucht, Probleme schnell identifiziert und teamorientiert angeht. Gerne beantworte ich Ihre Fragen in einem persönlichen Gespräch und freue mich auf die Einladung. 
''', style='bewerbung_default_paragraph')


document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

# document.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

saving_name = firma_receiver + ".docx"
document.save(saving_name)
print("Saved as ", saving_name)